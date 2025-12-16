import os
import PyPDF2
from docx import Document  

def extract_text_from_pdf(path):
    """Extract text from PDF safely."""
    text = ""
    try:
        if hasattr(path, "read"):
            reader = PyPDF2.PdfReader(path)
        else:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    except PyPDF2.errors.PdfReadError:
        print(f"⚠️ Cannot read PDF: {path}")
    except Exception as e:
        print(f"⚠️ Error reading PDF {path}: {e}")
    return text

def extract_text_from_docx(path):
    """Extract text from DOCX safely."""
    text = ""
    try:
        if hasattr(path, "read"):
            from io import BytesIO
            doc = Document(BytesIO(path.read()))
        else:
            doc = Document(path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"⚠️ Cannot read DOCX {path}: {e}")
    return text

def match_score(resume_text, jd_text):
    """Simple keyword match score."""
    score = 0
    jd_keywords = jd_text.lower().split()
    resume_words = resume_text.lower().split()
    matched_skills = []
    for word in jd_keywords:
        if word in resume_words:
            score += 1
            matched_skills.append(word)
    return score, matched_skills

def extract_skills_from_jd(jd_text):
    """
    Extract a simple list of skills/keywords from JD.
    You can improve this using NLP or a predefined skills list.
    """
    # For simplicity, split by spaces and remove duplicates
    skills = list(set(jd_text.lower().split()))
    return skills

def rank_candidates(resume_folder, jd_input, jd_is_text=False):
    """
    Rank candidates based on match scores and matched skills.

    resume_folder: folder containing resumes (PDF/DOCX)
    jd_input: job description path (str) or text (str)
    jd_is_text: True if jd_input is already text
    """
    # Get JD text
    if jd_is_text:
        jd_text = jd_input.lower()
    else:
        if jd_input.lower().endswith(".pdf"):
            jd_text = extract_text_from_pdf(jd_input).lower()
        elif jd_input.lower().endswith(".docx"):
            jd_text = extract_text_from_docx(jd_input).lower()
        else:
            raise ValueError("JD must be a PDF or DOCX file.")

    jd_skills = extract_skills_from_jd(jd_text)

    scores_dict = {}

    for resume_file in os.listdir(resume_folder):
        resume_path = os.path.join(resume_folder, resume_file)
        if resume_file.lower().endswith(".pdf"):
            resume_text = extract_text_from_pdf(resume_path).lower()
        elif resume_file.lower().endswith(".docx"):
            resume_text = extract_text_from_docx(resume_path).lower()
        else:
            print(f"⚠️ Skipping unsupported file type: {resume_file}")
            continue

        if resume_text.strip() == "":
            print(f"⚠️ Skipping empty/broken resume: {resume_file}")
            continue

        score, matched_skills = match_score(resume_text, jd_text)

        # Filter matched skills to only those in JD keywords
        matched_skills = [skill for skill in jd_skills if skill in resume_text]

        scores_dict[resume_file] = {
            "score": score,
            "matched_skills": matched_skills
        }

    ranked_list = sorted(
        scores_dict.items(),
        key=lambda x: x[1]["score"],
        reverse=True
    )

    return ranked_list
