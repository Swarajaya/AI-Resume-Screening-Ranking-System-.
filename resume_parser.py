from docx import Document
import os

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""

    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + " "

    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + " "

    return text.strip()

def parse_all_resumes(resume_folder):
    resumes = {}
    for file_name in os.listdir(resume_folder):
        if file_name.endswith(".docx"):
            file_path = os.path.join(resume_folder, file_name)
            resumes[file_name] = extract_text_from_docx(file_path)
    return resumes


if __name__ == "__main__":
    resume_texts = parse_all_resumes("data/resumes")
    for name, text in resume_texts.items():
        print(f"\n===== {name} =====\n")
        print(text[:1000])
import os

print("Files found in resumes folder:")
for f in os.listdir("data/resumes"):
    print(f)
